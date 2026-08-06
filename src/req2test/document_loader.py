"""Requirement document readers for text, Markdown, DOCX and PDF."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document
from pypdf import PdfReader


SUPPORTED_SUFFIXES = {".txt", ".md", ".docx", ".pdf"}


def _clean_text(text: str) -> str:
    lines = [line.rstrip() for line in text.replace("\r\n", "\n").replace("\r", "\n").split("\n")]
    cleaned: list[str] = []
    previous_blank = False
    for line in lines:
        is_blank = not line.strip()
        if is_blank and previous_blank:
            continue
        cleaned.append(line)
        previous_blank = is_blank
    return "\n".join(cleaned).strip()


def load_document(path: str | Path) -> str:
    file_path = Path(path)
    suffix = file_path.suffix.lower()
    if suffix not in SUPPORTED_SUFFIXES:
        raise ValueError(f"暂不支持文件类型：{suffix}")
    return load_document_bytes(file_path.read_bytes(), suffix)


def load_document_bytes(content: bytes, suffix: str) -> str:
    suffix = suffix.lower()
    if not suffix.startswith("."):
        suffix = f".{suffix}"

    if suffix in {".txt", ".md"}:
        for encoding in ("utf-8-sig", "utf-8", "gb18030"):
            try:
                return _clean_text(content.decode(encoding))
            except UnicodeDecodeError:
                continue
        raise ValueError("文本文件编码无法识别，请转换为 UTF-8 后重试")

    if suffix == ".docx":
        document = Document(BytesIO(content))
        blocks: list[str] = []
        blocks.extend(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    blocks.append(" | ".join(cells))
        return _clean_text("\n".join(blocks))

    if suffix == ".pdf":
        reader = PdfReader(BytesIO(content))
        pages = [(page.extract_text() or "").strip() for page in reader.pages]
        text = _clean_text("\n\n".join(page for page in pages if page))
        if not text:
            raise ValueError("PDF 未提取到可复制文本，扫描版 PDF 需要先进行 OCR")
        return text

    raise ValueError(f"暂不支持文件类型：{suffix}")
